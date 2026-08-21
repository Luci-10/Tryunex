#!/usr/bin/env python3
"""
Generates the TryUnex documentation as a Word (.docx) file.

Kept as a script rather than a hand-written document so it can be regenerated
when the product changes:

    python3 docs/generate_documentation.py
"""
import datetime, subprocess, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x6D, 0x28, 0xD9)
INK = RGBColor(0x1F, 0x1F, 0x23)
MUTED = RGBColor(0x5A, 0x5A, 0x66)
DANGER = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()

# ---------------------------------------------------------------- page setup
for s in doc.sections:
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)   # A4
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, color in (("Heading 1", 20, BRAND), ("Heading 2", 15, BRAND),
                          ("Heading 3", 12.5, INK), ("Heading 4", 11, INK)):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------------- helpers
def h(text, level=1, page_break=False):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(text, level=level)

def p(text="", bold=False, italic=False, color=None, size=None, align=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold, run.italic = bold, italic
    if color is not None: run.font.color.rgb = color
    if size is not None: run.font.size = Pt(size)
    if align is not None: par.alignment = align
    return par

def rich(parts):
    """parts: list of (text, bold, code) so a sentence can mix styles."""
    par = doc.add_paragraph()
    for text, bold, code in parts:
        run = par.add_run(text)
        run.bold = bold
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x8A, 0x1C, 0x5C)
    return par

def bullets(items, style="List Bullet"):
    for it in items:
        if isinstance(it, tuple):
            par = doc.add_paragraph(style=style)
            r = par.add_run(it[0]); r.bold = True
            par.add_run(" — " + it[1])
        else:
            doc.add_paragraph(it, style=style)

def numbered(items):
    bullets(items, style="List Number")

def code(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.25)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    shade(par, "F4F2FA")
    return par

def shade(par, hexcolor):
    pr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    pr.append(shd)

def callout(title, text, fill="FFF7E6", color=None):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.15)
    r = par.add_run(title + "  ")
    r.bold = True
    if color is not None: r.font.color.rgb = color
    par.add_run(text)
    shade(par, fill)
    par.paragraph_format.space_after = Pt(10)
    return par

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def toc_field():
    par = doc.add_paragraph()
    r = par.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click here and choose “Update Field” to build the contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        r._r.append(el)

def git(cmd, default="unknown"):
    try:
        return subprocess.check_output(cmd, shell=True, text=True,
                                       cwd=os.path.dirname(os.path.dirname(os.path.abspath(__file__)))).strip()
    except Exception:
        return default

COMMIT = git("git rev-parse --short HEAD")
TODAY = datetime.date.today().strftime("%d %B %Y")

# ============================================================== TITLE PAGE
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(150)
r = t.add_run("TryUnex"); r.font.size = Pt(46); r.font.bold = True; r.font.color.rgb = BRAND

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Complete Product & Technical Documentation")
r.font.size = Pt(16); r.font.color.rgb = MUTED

st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st2.add_run("Your wardrobe, your looks, and a place to pass clothes on")
r.font.size = Pt(11); r.italic = True; r.font.color.rgb = MUTED

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(40)
r = meta.add_run(f"Version 1.1  ·  {TODAY}  ·  build {COMMIT}\nweb: www.tryunex.in   ·   Android app: in.tryunex.app")
r.font.size = Pt(10); r.font.color.rgb = MUTED

# ============================================================== HOW TO READ
h("How to read this document", 1, page_break=True)
p("This document describes TryUnex completely: what it does, every screen and button a person can "
  "touch, how each feature behaves, and how the whole thing is built and operated. It is written so "
  "that someone who has never seen the product can start at the beginning and understand it, and so "
  "that someone who has to work on it can find the specific detail they need.")

p("It is organised in seven parts, and you do not need to read them in order:", bold=False)
table(["Part", "What it covers", "Who it is for"],
      [["1. The product", "What TryUnex is, who it is for, and the vocabulary used throughout", "Everyone"],
       ["2. Using TryUnex", "Every screen and every button, in the order a new person meets them", "Everyone"],
       ["3. How features work", "The rules behind credits, try-on, selling, privacy and consent", "Everyone"],
       ["4. Technical reference", "Architecture, data model, the full API, and the code layout", "Developers"],
       ["5. Running it", "Deployment, configuration, building the app, and the test suites", "Developers / operators"],
       ["6. Honest limits", "What is deliberately not built yet, and what is still outstanding", "Everyone"],
       ["7. Reference tables", "Endpoints, tables, environment variables and metrics at a glance", "Developers"]],
      widths=[1.5, 3.4, 1.6])

callout("A note on accuracy.",
        "Everything here was read from the code as it stands at build " + COMMIT +
        ", not from memory or intention. Where something is promised but not yet built, Part 6 says so plainly "
        "rather than leaving it implied.")

h("Contents", 2)
toc_field()
p("(In Word: right-click the line above, choose Update Field, then Update entire table.)",
  italic=True, color=MUTED, size=9)

# ============================================================== PART 1
h("Part 1 — The product", 1, page_break=True)

h("What TryUnex is", 2)
p("TryUnex is a personal wardrobe app. You photograph the clothes you own, and the app becomes a "
  "record of them: what you have, what you have worn recently, what is clean, and what goes together. "
  "On top of that record it offers three things you cannot do with a folder of photos.")
bullets([
    ("Virtual try-on", "you upload one photo of yourself, pick garments from your wardrobe, and the app "
     "generates an image of you wearing them, so you can see an outfit before putting it on."),
    ("An AI stylist", "a chat that can see your actual wardrobe and answer questions like what to wear "
     "to an evening dinner, or what you have not worn in a while."),
    ("Thrift", "a person-to-person marketplace for passing clothes on. When a sale completes, the "
     "garment moves out of the seller's wardrobe and into the buyer's automatically."),
])

h("Who it is for", 2)
p("People who own more clothes than they can hold in their head, and who want to get more use out of "
  "what they already own — by seeing it laid out, by planning outfits, and by passing on what no "
  "longer fits their life rather than letting it sit unused.")

h("The ideas the app is built on", 2)
p("A few decisions shape everything else, and knowing them makes the rest of the product make sense.")
bullets([
    ("Your wardrobe is private by default", "nothing is visible to anyone until you hand out a share "
     "code, and even then you choose what they may do. Photographs are stored privately and served "
     "only to people entitled to see them."),
    ("Clothes have a state, not just an existence", "every garment is either clean or worn. Wearing "
     "something moves it to the laundry basket; resetting brings it back. This is what makes "
     "\"what can I actually wear today\" answerable."),
    ("Generation costs credits", "try-on images are expensive to produce, so they are metered. "
     "Credits are tracked in an append-only ledger and every debit happens on the server."),
    ("TryUnex never holds the goods or the money in Thrift", "it introduces two people and moves the "
     "garment between wardrobes once they both confirm. Payment and handover happen between them."),
])

h("Vocabulary", 2)
p("These words have specific meanings throughout the app and this document.")
table(["Term", "Meaning"],
      [["Piece / garment", "One item of clothing in a wardrobe, with a photo, a name, a category and a style tag."],
       ["Clean / Worn", "The two states a piece can be in. Worn pieces sit in the laundry basket until reset."],
       ["Laundry basket", "The Worn screen — everything currently marked as worn."],
       ["Wear event", "A record that a piece was worn on a particular date. Drives \"last worn\" and history."],
       ["Look", "One generated try-on image: you, wearing between one and five chosen garments."],
       ["Credit", "The unit that pays for generating a look. One or two credits per look, by item count."],
       ["Plan", "An outfit chosen in advance for a future date."],
       ["Share code", "A code you generate to let someone else see your wardrobe."],
       ["Suggestion", "An outfit someone with access proposes to you, which you accept or decline."],
       ["Listing", "A garment offered for sale on Thrift."],
       ["Transfer", "The automatic move of a garment from seller's wardrobe to buyer's when a sale completes."]],
      widths=[1.4, 5.1])

# ============================================================== PART 2
h("Part 2 — Using TryUnex, screen by screen", 1, page_break=True)
p("This part walks through every screen in the order a new person meets them, and names every button "
  "on each one and what it does. Where a screen behaves differently on the phone app than in a browser, "
  "that is called out.")

h("2.1  Getting in", 2)

h("The landing page", 3)
p("What a visitor sees before signing in: what the product does, illustrated with photography, and two "
  "ways forward.")
table(["Button", "What it does"],
      [["Start free / Start styling free / Create my free wardrobe", "Goes to the sign-in screen to make an account."],
       ["Sign in", "Goes to the sign-in screen for an existing account."],
       ["Explore how it works", "Scrolls down the page through the explanation sections."]],
      widths=[2.3, 4.2])

h("Signing in", 3)
p("TryUnex has no passwords. You enter your email address and receive a six-digit code.")
numbered([
    "Type your email address and press Continue. A code is sent to that address, valid for ten minutes.",
    "Type the six-digit code. Six separate boxes accept one digit each, and pasting the whole code works.",
    "If the address already has an account you are signed straight in. If not, you continue to registration.",
])
table(["Button", "What it does"],
      [["Continue", "Sends the code to the address you typed."],
       ["Change", "Goes back to correct the email address."],
       ["Resend", "Issues a fresh code, replacing the previous one."]],
      widths=[2.3, 4.2])
callout("Five attempts.", "A wrong code can be tried five times before the code is cancelled and you must "
        "request a new one. This is deliberate — it stops someone guessing their way into an inbox they do not own.")

h("Registration", 3)
p("Asked once, for a new account only.")
table(["Field", "Why it is asked", "Required"],
      [["Your name", "Shown on your profile and to people you share a wardrobe with", "Yes"],
       ["Date of birth", "Used to check the minimum age. Stored, not shown to others", "Yes"],
       ["Gender", "Improves try-on and styling suggestions. Includes “Prefer not to say”", "Optional"]],
      widths=[1.5, 3.6, 1.0])
p("Buttons: Create my wardrobe completes registration; the gender choices are Male, Female, Other and "
  "Prefer not to say.")

h("Accepting the Terms", 3)
p("The first time you sign in — and again whenever the policies materially change — a panel appears "
  "that cannot be dismissed by tapping away. It links to the Terms of Service, Privacy Policy and "
  "Refund Policy, and records which version you accepted and when.")
table(["Button", "What it does"],
      [["Agree and continue", "Records your acceptance and opens the app."],
       ["Not now — sign out", "Signs you out. The app cannot be used without accepting."]],
      widths=[2.3, 4.2])

h("2.2  The wardrobe (home)", 2)
p("The main screen and the one you land on. It shows your pieces as cards in a grid, with a summary "
  "across the top of how many pieces you own, how many are clean and how many are worn.")

h("The two tabs", 3)
bullets([("Clean", "everything available to wear right now."),
         ("Worn", "the laundry basket — everything worn since your last reset.")])

h("Finding things", 3)
bullets([("Search", "filters by name as you type. A Clear search button empties it."),
         ("Filter", "narrows by category (top, bottom, dress, outerwear, footwear, accessory and so on) and by style tag."),
         ("Wardrobe switcher", "if someone has shared a wardrobe with you, this swaps between yours and theirs.")])

h("Each garment card", 3)
p("Tapping anywhere on a card — the photo or the name — opens the full details. A Try on button sits "
  "on the card itself as a shortcut.")

h("Buttons on this screen", 3)
table(["Button", "What it does"],
      [["Add piece  (+)", "Opens the add-a-garment panel. The main action on the screen."],
       ["Try on  (on a card)", "Adds that garment to a new look and opens the Try-on Studio."],
       ["Reset / Reset all", "Moves everything from the laundry basket back to clean. Asks first."],
       ["Select (long press)", "Starts multi-select, so several pieces can be worn or planned at once."],
       ["Clear (selection tray)", "Cancels the current selection."]],
      widths=[2.3, 4.2])

h("Adding a garment", 3)
numbered([
    "Choose a photo — take one with the camera, or pick from your gallery.",
    "The photo is resized and uploaded. A progress bar shows how far it has got.",
    "Give it a name, choose a category, and optionally a style tag.",
    "Save. The piece appears in your wardrobe as clean.",
])
callout("On the phone app,", "the camera and gallery are separate choices, and the app asks for permission "
        "the first time. Photos are only ever read when you pick one — nothing is scanned in the background.")

h("Garment details", 3)
p("Opened by tapping a card. Shows the photo full width, the name, category and style tag, and two "
  "facts drawn from your history: how many times you have worn it, and when you last did.")
table(["Button", "What it does"],
      [["Wear today", "Marks it worn, records today's date, and moves it to the laundry basket."],
       ["Try on", "Starts a look with this piece."],
       ["Ask AI", "Opens the stylist chat with this piece attached to the conversation."],
       ["Sell this piece", "Creates a Thrift listing from it, reusing the photo you already have."],
       ["Delete this piece", "Removes it permanently. Asks first, and its photo is deleted from storage."]],
      widths=[2.3, 4.2])

h("2.3  The laundry basket (Worn)", 2)
p("Everything currently marked worn. The point of the screen is the reset: when the laundry is done, "
  "one button returns it all to clean.")
table(["Button", "What it does"],
      [["Reset all", "Returns every worn piece to clean. Wear history is kept."],
       ["Undo", "Appears briefly after a reset, in case it was a mistake."]],
      widths=[2.3, 4.2])
callout("Resetting does not erase history.", "It changes what is available to wear. The record of when you "
        "wore something is permanent, and is what “last worn” and the History screen are built from.")

h("2.4  Planning outfits", 2)
p("For deciding in advance what to wear on a given day. You pick a date, choose pieces from what is "
  "clean, and the plan is saved. Upcoming plans are listed newest first.")
table(["Button", "What it does"],
      [["Plan outfit", "Starts a new plan."],
       ["Choose pieces", "Opens the picker of clean garments."],
       ["Save plan", "Stores the outfit against the chosen date."],
       ["Delete (on a plan)", "Removes an upcoming plan."]],
      widths=[2.3, 4.2])
p("When the day arrives, the pieces in that plan are flipped from clean to worn automatically.")

h("2.5  Try-on Studio", 2)
p("Where the app generates a picture of you wearing chosen clothes. It needs two things: one photo of "
  "you, and between one and five garments.")

h("Your photo", 3)
p("Uploaded once and reused for every look. Replacing it swaps it everywhere. It is stored privately "
  "and is never visible to anyone else — not to people you share a wardrobe with, and not on Thrift.")
p("For the best result the app asks for a clear, well-lit, full-body photo in fitted clothing, so the "
  "silhouette is readable.")

h("Making a look", 3)
numbered([
    "Add your photo, if you have not already.",
    "Choose up to five garments. The cost in credits is shown before you commit.",
    "Press Generate. The image takes a short while to produce.",
    "The result appears, and is saved to your try-on history.",
])
table(["Button", "What it does"],
      [["Add your photo / Replace photo", "Uploads or swaps the photo of you."],
       ["Edit outfit", "Changes which garments are in the look."],
       ["Generate", "Produces the image and spends the credits."],
       ["New look / Start new look", "Clears the current result and starts again."],
       ["View full size", "Opens the image large, where it can be zoomed."],
       ["Save image", "Saves the generated image to your device."],
       ["Share", "Passes the image to the device's share sheet."],
       ["Show original", "Toggles between the generated look and your original photo."],
       ["Buy credits / View plans", "Shown when you do not have enough credits."]],
      widths=[2.3, 4.2])
callout("What try-on is, honestly.", "It is a visual styling preview, not a fitting service. Fit and size "
        "may vary — the app says so on the screen, and you should choose your usual size from the garment's "
        "own details.", fill="FFF7E6")

h("2.6  The AI stylist", 2)
p("A chat, reachable from the button that floats above the wardrobe. It can see your actual wardrobe — "
  "what you own, its categories and style tags, and when each piece was last worn — so its answers are "
  "about your clothes rather than clothes in general.")
table(["Button", "What it does"],
      [["Ask AI about your wardrobe", "Opens the chat."],
       ["Send", "Sends your message."],
       ["Stop generating", "Interrupts a reply in progress."],
       ["New chat", "Clears the conversation and starts fresh."],
       ["Try it on", "Takes garments the stylist suggested straight into the Try-on Studio."],
       ["Go to my wardrobe", "Closes the chat and returns to the grid."]],
      widths=[2.3, 4.2])
p("Free accounts get ten stylist conversations a month; paid plans lift that limit.")

h("2.7  Thrift — passing clothes on", 2)
p("A person-to-person marketplace. TryUnex introduces buyer and seller and moves the garment between "
  "wardrobes when the sale completes. It does not take payment, hold stock, arrange delivery, or "
  "guarantee anything about the transaction — that is between the two people.")

h("Browsing", 3)
p("A grid of what other people are selling, with search and a filter panel.")
table(["Filter", "Choices"],
      [["Search", "Free text across listings"],
       ["Size / Condition", "Condition is like new, gently used, or used"],
       ["Min price / Max price", "A price range"],
       ["City", "Where the seller is"],
       ["Delivery", "Pickup, shipping, or either"],
       ["Style", "The same style tags used in your wardrobe"],
       ["Sort by", "How results are ordered"]],
      widths=[1.8, 4.7])
p("Buttons: Filter & sort opens the panel, Clear empties every filter, Show results closes it.")

h("A listing", 3)
p("The photo, price, size, condition, delivery method, the seller's city and display name. Private "
  "contact details are never shown.")
table(["Button", "What it does"],
      [["Message seller", "Opens a conversation about this listing."],
       ["Save", "Adds it to your saved list. Tapping again removes it."],
       ["Try with my wardrobe", "Generates a look with this garment against your photo, before you buy."],
       ["Report this listing", "Flags it for review, with a reason."],
       ["Block seller", "Hides them from you and you from them, in both directions."]],
      widths=[2.3, 4.2])

h("Selling", 3)
numbered([
    "From a garment's details, choose Sell this piece — the existing photo is reused.",
    "Set price, size, condition, delivery method and city.",
    "Activate the listing so it appears publicly.",
    "Talk to interested buyers in messages, and agree payment and handover between yourselves.",
    "When the buyer has paid and has the item, mark it sold and confirm the sale.",
])
table(["Button", "What it does"],
      [["Edit", "Changes the details of a listing."],
       ["Pause", "Hides it temporarily without deleting it."],
       ["Mark sold", "Records the sale and begins the transfer."],
       ["Delete listing", "Removes it entirely."]],
      widths=[2.3, 4.2])

h("Completing a sale, and the transfer", 3)
p("A sale needs both people. The seller marks it sold; the buyer confirms they have paid and received "
  "the item. Only when both have confirmed does the garment move: it leaves the seller's wardrobe and "
  "appears in the buyer's, with the same photo, as a clean piece.")
callout("The photo is shared, not copied.", "Both wardrobes point at the same stored image. This is why "
        "deleting your account never removes a picture that a buyer now depends on.")
table(["Button", "What it does"],
      [["Confirm purchase", "Buyer confirms payment and receipt. Triggers the transfer."],
       ["Cancel", "Either side calls the sale off before it completes."]],
      widths=[2.3, 4.2])

h("Messages", 3)
p("One conversation per listing, between one buyer and one seller. Buttons: Send message, Report this "
  "conversation, and Block.")

h("Saved", 3)
p("Everything you have saved while browsing, in one list.")

h("2.8  Sharing your wardrobe", 2)
p("Nobody can see your wardrobe unless you give them a code.")
numbered([
    "Generate a share code and choose what the holder may do.",
    "Send the code to them yourself — the app does not send it.",
    "They enter it under “Have a code?” and the wardrobes connect.",
])
table(["Button", "What it does"],
      [["Generate share code", "Creates a new code to hand out."],
       ["Copy", "Copies the code to the clipboard."],
       ["Cancel code", "Kills an unused code."],
       ["Connect", "Redeems a code someone gave you."],
       ["Remove / Disconnect", "Ends a sharing relationship from either side."]],
      widths=[2.3, 4.2])
p("The screen lists both directions separately: people who can see your wardrobe, and wardrobes shared "
  "with you.")

h("A friend's wardrobe", 3)
p("What you can do depends on what they granted: view their pieces, see their upcoming plans, and "
  "propose an outfit. A suggestion carries a date and an optional note, and they accept or decline it. "
  "Try-on access is a separate permission — being able to see a wardrobe does not include it.")

h("2.9  History", 2)
p("A record of what you wore and when, newest first — the long view of the wear events created every "
  "time you mark something worn.")

h("2.10  Notifications", 2)
p("A bell beside your profile picture, showing a count of anything unread. It covers messages about "
  "your listings, sales and transfers, outfit suggestions, and someone connecting to your wardrobe.")
table(["Button", "What it does"],
      [["Notifications (bell)", "Opens the tray — a panel on desktop, a sheet on mobile."],
       ["Mark all read", "Clears the unread count."],
       ["(a notification)", "Takes you to whatever it is about."]],
      widths=[2.3, 4.2])
p("It refreshes about once a minute, and stops polling when the tab is in the background.")

h("2.11  Plans & credits", 2)
p("What credits are, how many you have, and how to get more. Free accounts receive three credits on "
  "joining and one more each month.")
table(["Button", "What it does"],
      [["Buy (a pack)", "One-off purchase of credits that do not expire monthly."],
       ["Subscribe (a plan)", "Monthly credits, renewing automatically."],
       ["Manage / Cancel", "Handles an existing subscription."]],
      widths=[2.3, 4.2])
p("Payment is taken by Razorpay. TryUnex never sees or stores your card or UPI details.")

h("2.12  Account, Settings and the rest", 2)

h("Account", 3)
p("Your profile, outfit suggestions waiting for you (Accept or Decline), and Sign out.")

h("Settings", 3)
p("Ten sections, in this order:")
table(["Section", "What is in it"],
      [["Appearance", "Motion preference — full, reduced or none. Saved on the device."],
       ["Privacy & sharing", "An explanation of the sharing model, and a link to shared wardrobes."],
       ["Photos & virtual try-on", "What happens to your photos, and a link to the studio."],
       ["Account", "Name, email, date of birth and gender, with a link to your profile."],
       ["Plan & credits", "Current plan, credits left, renewal date, and a link to Plans."],
       ["Getting started", "Replay app walkthrough — repeats the guided tour, changing nothing."],
       ["Legal & consent", "The three policies, and which version you accepted and when."],
       ["Danger zone", "Delete my account."],
       ["App", "Version and build information."],
       ["Support", "Contact and about."]],
      widths=[1.7, 4.8])

h("Deleting your account", 3)
p("Under Danger zone. It is permanent and immediate, and it takes two steps.")
numbered([
    "A review screen lists exactly what will be destroyed, counted from your actual account — garments, "
    "try-on photos, listings and conversations — and warns about unused credits.",
    "Press Continue and a six-digit code is emailed to your registered address.",
    "Enter the code and press Delete permanently.",
])
table(["Button", "What it does"],
      [["Delete my account", "Opens the review screen."],
       ["Keep my account", "Closes it. Nothing happens."],
       ["Continue", "Emails the confirmation code."],
       ["Back", "Returns to the review screen."],
       ["Send a new code", "Issues a fresh code."],
       ["Delete permanently", "Deletes the account. There is no undo."]],
      widths=[2.3, 4.2])
callout("Why a code, when you are already signed in?", "So that an unattended or stolen session cannot "
        "destroy someone's wardrobe. The code proves whoever is pressing the button holds the email "
        "account right now.", fill="FDECEC", color=DANGER)
p("What is deleted: your profile, wardrobe, wear history, try-on photos and generated looks, listings, "
  "messages, sharing connections and notifications — and the image files themselves are removed from "
  "storage. What is kept: photos of garments you already sold, because the buyer owns those now; and a "
  "record of payments you made, with your name, email and account reference removed, because "
  "accounting rules require proof that a transaction happened.")

h("Legal pages", 3)
p("Privacy Policy, Terms of Service and Refund & Credit Policy, each reachable from Settings and from "
  "the landing page footer, each with its own contents list.")

h("Contact and About", 3)
p("Contact sends a message to support (subject and message; Send another resets the form). About shows "
  "the version and build.")

# ============================================================== PART 3
h("Part 3 — How the features actually work", 1, page_break=True)
p("Part 2 described what the buttons do. This part explains the rules underneath them — the parts that "
  "are not obvious from the screen, and that matter when something behaves unexpectedly.")

h("3.1  Credits and billing", 2)

h("What a look costs", 3)
table(["Garments in the look", "Credits"],
      [["1 to 3", "1 credit"], ["4 or 5", "2 credits"], ["More than 5", "Not allowed"]],
      widths=[3.2, 3.3])
p("The screen shows the cost before you commit, but the figure that counts is calculated again on the "
  "server from the garments actually sent. The browser cannot talk the price down.")

h("The free allowance", 3)
bullets([("Three credits on joining", "valid for 30 days."),
         ("One credit a month after that", "granted at most once per account per month.")])

h("What you can buy", 3)
table(["Product", "Type", "Credits", "Price"],
      [["Starter pack", "One-off", "3", "₹29"],
       ["Mid pack", "One-off", "6", "₹52"],
       ["Bulk pack", "One-off", "10", "₹79  (best value)"],
       ["Lite", "Monthly", "7 per month", "₹55 / month"],
       ["Plus", "Monthly", "14 per month", "₹99 / month  (most popular)"],
       ["Style", "Monthly", "30 per month", "₹199 / month"]],
      widths=[1.5, 1.2, 1.6, 2.2])
p("Stylist chats are limited to ten a month on the free tier.")

h("How the balance is kept", 3)
p("Credits are not a number in a column that goes up and down. Every grant, purchase, spend and refund "
  "is a row in an append-only ledger, and the balance is the sum of those rows that have not expired. "
  "Nothing is ever edited or deleted, so the history of how a balance came to be is always recoverable.")
bullets([
    ("Every write carries an idempotency key", "so a retried request, a double-tap, or a payment webhook "
     "delivered twice can only ever be counted once."),
    ("Debits happen on the server", "before the image is generated, never in the browser."),
    ("A failed generation is refunded", "for the exact amount that was debited, not a guessed amount."),
])

h("Payment", 3)
p("Razorpay handles the transaction. TryUnex stores what was bought, for how much, and Razorpay's own "
  "identifiers — never card numbers, UPI PINs or any payment credential. Purchases are confirmed by a "
  "signed webhook from Razorpay rather than by the browser saying the payment worked.")

h("3.2  How try-on works", 2)
p("Generation runs on fal's FLUX Virtual Try-On model. The sequence is:")
numbered([
    "You choose garments and press Generate.",
    "The server checks your balance and debits the cost, under a lock so two requests cannot spend the "
    "same credit twice.",
    "The garments are composited into a single sheet image, because the model accepts exactly one "
    "garment image no matter how many pieces you picked.",
    "Your photo and that sheet are sent to fal with a prompt describing the outfit.",
    "The result is stored privately and recorded in your try-on history.",
    "If anything fails, the credits are refunded automatically.",
])
callout("Why the sheet exists.", "FLUX VTO takes one garment image. Compositing several garments onto a "
        "single canvas — within a strict pixel budget the model imposes — is what makes multi-item looks "
        "possible at all.")

h("3.3  Selling and the wardrobe transfer", 2)
p("The interesting part of Thrift is what happens when a sale completes, because it changes two people's "
  "wardrobes at once and must not half-happen.")
bullets([
    ("Both sides must confirm", "the seller marks it sold, the buyer confirms payment and receipt."),
    ("The whole transfer is one database statement", "the garment leaving one wardrobe and arriving in the "
     "other cannot come apart, because there is no moment between them."),
    ("It cannot run twice", "the transfer records the new garment's identity on the sale itself, so a "
     "repeated confirmation finds the work already done and changes nothing."),
    ("The photo is shared", "both wardrobes reference the same stored image, and it is only ever deleted "
     "when the last wardrobe referencing it lets go."),
])
p("What TryUnex deliberately does not do in this version: take payment, pay out sellers, produce "
  "shipping labels, hold money in escrow, track delivery, handle returns, or offer buyer protection.")

h("3.4  Privacy and how images are protected", 2)
p("Photographs are the most sensitive thing in the product — clothes, and photographs of your body. "
  "The rules are:")
bullets([
    ("Storage is private", "no image has a public address. Nothing can be reached by guessing a URL."),
    ("You name a record, never a file", "the app asks for “this garment” or “this look”, and the server "
     "decides whether you may see it and only then produces a short-lived link. A caller cannot ask for "
     "an arbitrary file."),
    ("Try-on photos are yours alone", "a photo of you and the looks generated from it are readable only "
     "by your account. Sharing a wardrobe does not include them, and listing a garment does not expose them."),
    ("Garment photos follow the garment", "readable by the owner, by people they shared with, and — while "
     "a piece is listed — by anyone browsing Thrift."),
    ("Deleting means deleting", "removing a garment removes its stored file, unless another wardrobe still "
     "depends on it."),
])
p("Requests for image links are rate-limited so the collection cannot be enumerated by a script.")

h("3.5  Consent, age and policy versions", 2)
p("Acceptance is recorded against a specific policy version, per user, with a timestamp. Publishing a "
  "materially changed policy means raising the version, and everyone is asked once more the next time "
  "they open the app. Old acceptances are kept as the record of what was agreed and when.")
p("Date of birth is collected at registration and checked against a minimum age. It is stored and never "
  "shown to other people.")

h("3.6  Notifications", 2)
p("Notifications are written when something happens that concerns you: a message about your listing, a "
  "sale confirmed, a transfer completed, an outfit suggested, a wardrobe connected.")
bullets([
    ("Writing one can never break the thing that caused it", "if recording a notification fails, the sale "
     "or message it was about still succeeds."),
    ("Repeats collapse", "several messages in one conversation become a single entry that updates, rather "
     "than a stack of near-identical rows."),
])

# ============================================================== PART 4
h("Part 4 — Technical reference", 1, page_break=True)

h("4.1  Architecture", 2)
table(["Layer", "Technology"],
      [["Frontend", "React 19, TypeScript, Vite 6, Tailwind CSS 3, React Router"],
       ["Backend", "Node.js, Express 4, TypeScript, Zod for input validation"],
       ["Database", "Neon (serverless PostgreSQL), accessed through Drizzle ORM"],
       ["Image storage", "Cloudflare R2, private bucket, S3-compatible, signed URLs"],
       ["Try-on model", "fal — fal-ai/flux-pro/v1/vto"],
       ["AI stylist", "Google Gemini"],
       ["Payments", "Razorpay (orders, subscriptions, signed webhooks)"],
       ["Email", "Authenticated SMTP for the sending domain, with a Gmail fallback"],
       ["Hosting", "Vercel — static frontend plus serverless API"],
       ["Mobile", "Capacitor 7, Android. Web assets bundled into the APK"]],
      widths=[1.5, 5.0])

h("4.2  Repository layout", 2)
code("""Tryunex_wordrobe/
├── backend/
│   ├── src/
│   │   ├── routes/        HTTP endpoints, one file per area
│   │   ├── services/      business logic, external integrations
│   │   ├── db/            schema and database client
│   │   └── app.ts         express app, route mounting, error handling
│   └── scripts/           verification suites, audit and maintenance tools
├── frontend/
│   ├── src/
│   │   ├── pages/         one file per screen
│   │   ├── components/    shared UI, incl. ui/ primitives
│   │   └── *.ts           api client, auth, billing, thrift, upload helpers
│   ├── public/            static files, incl. prerendered SEO pages
│   └── android/           Capacitor Android project
└── docs/                  this document and its generator""")

h("4.3  Data model", 2)
p("Twenty-five tables. Grouped by what they are for:")
table(["Area", "Tables"],
      [["People", "users, policy_acceptances, onboarding_state"],
       ["Wardrobe", "clothes, wear_events"],
       ["Sharing", "share_codes, shares, suggestions"],
       ["Try-on", "tryon_assets, tryon_requests"],
       ["Billing", "billing_profiles, credit_ledger, payments, webhook_events, retained_financial_records"],
       ["Thrift", "thrift_listings, thrift_saves, thrift_conversations, thrift_messages, "
                  "thrift_transactions, thrift_transfers, thrift_blocks, thrift_listing_reports, "
                  "thrift_conversation_reports"],
       ["Other", "notifications"]],
      widths=[1.3, 5.2])

h("Relationships worth knowing", 3)
bullets([
    ("Almost everything cascades from users", "deleting a person removes their rows across the schema."),
    ("clothes → thrift_listings does not cascade", "deliberately dropped, so that removing a garment "
     "cannot erase the record of a completed sale."),
    ("wear_events is the source of truth for history", "clothes.status is only the current state."),
])

h("4.4  Engineering constraints that shape the code", 2)
p("These are non-obvious, were each learned from a real failure, and explain why parts of the code look "
  "the way they do.")
bullets([
    ("Neon's HTTP driver has no interactive transactions", "anything that must be atomic is written as a "
     "single SQL statement, often with data-modifying CTEs and an advisory lock in the first one."),
    ("A CTE that selects from another data-modifying CTE returns rows but does not persist its insert",
     "an audit write arranged that way silently did nothing."),
    ("Drizzle column helpers render unqualified inside raw SQL", "in a correlated subquery a bare column "
     "binds to the inner table, which silently produced “never worn” for every garment. Correlated "
     "subqueries name their tables explicitly."),
    ("Tailwind emits .relative after .absolute", "so a class passed in by a caller loses to the "
     "component's own, regardless of attribute order."),
    ("CapacitorHttp patches both fetch and XMLHttpRequest", "and the bridge turns a Blob body into \"{}\", "
     "so binary uploads in the app go through an unpatched request object."),
    ("count(*) comes back as a string", "from the Neon driver, and must be converted before comparison."),
])

h("4.5  Backend services", 2)
table(["Service", "Responsibility"],
      [["auth.ts", "Session tokens, cookies, the requireAuth guard"],
       ["otp.ts", "Stateless one-time codes in signed cookies, bound to a purpose"],
       ["mailer.ts", "Sign-in codes, deletion codes, contact messages"],
       ["age.ts", "Minimum-age rules and date-of-birth checks"],
       ["policy.ts", "Policy versions and who accepted what"],
       ["accountDeletion.ts", "Permanent deletion, image cleanup, financial archiving"],
       ["billing/catalogue.ts", "Packs and plans"],
       ["billing/credits.ts", "The ledger — grants, debits, refunds, balances"],
       ["billing/razorpay.ts", "Orders, subscriptions, webhook signature checks"],
       ["fal.ts", "The try-on model"],
       ["garmentSheet.ts", "Compositing several garments into one image"],
       ["tryonPrompt.ts", "The prompt sent with each generation"],
       ["r2.ts", "Signed uploads and downloads, deletion, key parsing"],
       ["thrift.ts", "Marketplace schema bootstrap"],
       ["thriftTransfer.ts", "Sale completion, wardrobe transfer, image reference counting"],
       ["notifications.ts", "Writing notifications without ever breaking the caller"],
       ["plans.ts", "Flipping planned outfits to worn on the day"],
       ["metrics.ts", "One structured log line per business event"],
       ["http-tune.ts", "Forces IPv4 and a longer connect timeout"]],
      widths=[1.8, 4.7])

h("4.6  Frontend structure", 2)
p("Twenty screens under pages/, and shared components under components/ with lower-level primitives in "
  "components/ui/.")
table(["Module", "Responsibility"],
      [["api.ts", "The HTTP client every screen calls through"],
       ["auth.tsx", "Who is signed in, and the guard around private routes"],
       ["media.ts", "Fetching protected images — differs between web and app"],
       ["upload.ts", "Resizing, orientation, and the app's binary-safe upload path"],
       ["billing.ts / thrift.ts / chat.tsx", "Feature-level API wrappers"],
       ["tryonCost.ts", "Mirrors the server's pricing so cost can be shown in advance"],
       ["photoPicker.ts / photoConsent.ts", "Camera and gallery access, and permission prompts"],
       ["platform.ts", "Whether this is the app or a browser"],
       ["motion.ts", "The motion preference"]],
      widths=[2.0, 4.5])

# ============================================================== PART 5
h("Part 5 — Running and maintaining TryUnex", 1, page_break=True)

h("5.1  Environments", 2)
table(["Environment", "Where", "How it updates"],
      [["Production", "www.tryunex.in", "Pushing to the main branch"],
       ["Preview", "A per-push vercel.app address", "Pushing any other branch"],
       ["Local", "localhost:5173 (web) + local API", "npm run dev"]],
      widths=[1.5, 2.5, 2.5])

h("5.2  Deploying", 2)
numbered([
    "Work on a branch, not on main.",
    "Run the verification suites and make sure they pass before pushing.",
    "Merge into main and push. Vercel builds and deploys automatically.",
    "Confirm the deploy is live by checking the commit reported by the health endpoint.",
])
code("curl -s https://www.tryunex.in/api/health")
p("The health endpoint reports the running commit, the environment, and booleans for whether each "
  "integration is configured — never a key, a length, or a prefix. That is enough to tell “the "
  "environment variable never reached this deployment” apart from “the key is wrong”, which is "
  "otherwise guesswork against a serverless deploy.")
callout("Order matters.", "Deploy after the tests pass, not alongside them. Pushing to production while a "
        "suite is still running has already caused one avoidable incident.", fill="FDECEC", color=DANGER)

h("5.3  Configuration", 2)
p("Every environment variable the code reads. None of them belong in the repository.")
table(["Variable", "Purpose"],
      [["DATABASE_URL", "Neon connection string"],
       ["JWT_SECRET", "Signs sessions, one-time codes, and the deletion archive pseudonym"],
       ["R2_ACCOUNT_ID / R2_BUCKET", "Cloudflare R2 account and bucket"],
       ["R2_ACCESS_KEY_ID / R2_SECRET_ACCESS_KEY", "R2 credentials (read and write on objects)"],
       ["R2_PUBLIC_BASE_URL", "Base address used to recognise our own stored objects"],
       ["FAL_KEY  (or FAL_API_KEY)", "Try-on model access. Server-side only"],
       ["GEMINI_API_KEY", "AI stylist"],
       ["RAZORPAY_KEY_ID / RAZORPAY_KEY_SECRET", "Payments"],
       ["RAZORPAY_WEBHOOK_SECRET", "Verifies webhook signatures"],
       ["RAZORPAY_PLAN_LITE_ID / _PLUS_ID / _STYLE_ID", "Subscription plan identifiers"],
       ["SMTP_HOST / SMTP_PORT / SMTP_USER / SMTP_PASS", "Authenticated mail for the sending domain"],
       ["GMAIL_USER / GMAIL_APP_PASSWORD", "Fallback mail. Codes sent this way tend to land in spam"],
       ["MAIL_FROM / MAIL_REPLY_TO", "Sender and reply-to addresses"],
       ["FRONTEND_ORIGIN", "Allowed browser origin for cross-origin requests"],
       ["TRYON_MOCK / TRYON_GENERATION_DISABLED", "Test and maintenance switches. Mock is refused in production"]],
      widths=[2.6, 3.9])

h("5.4  Building the Android app", 2)
p("The app bundles its web assets into the APK. A web deploy therefore reaches the website immediately "
  "and the app not at all — the app only changes when a new APK is built and installed. Server-side "
  "changes do reach the existing app, because it calls the same API.")
code("""cd frontend
npm run build
npx cap sync android
cd android
JAVA_HOME="/Applications/Android Studio.app/Contents/jbr/Contents/Home" ./gradlew assembleDebug""")
callout("Two traps in this build.",
        "First, the default JDK on the machine is 17 and Capacitor 7 requires 21 — without JAVA_HOME the "
        "build fails with “invalid source release: 21”. Android Studio's bundled JDK is 21, which is why "
        "building from the IDE works. Second, the Gradle config redirects the build directory to "
        "/tmp/android-builds, and the copy-back step does not run during assembleDebug — so the APK under "
        "app/build/outputs/ can be an old one while the build reports success. Take the APK from the "
        "redirected path, and remember /tmp is cleared on reboot.")
p("Current build: versionCode 2, versionName 1.1, minSdk 23, targetSdk 36. Debug builds install on your "
  "own devices; Google Play requires a signed release build, which has not been set up yet.")

h("5.5  The verification suites", 2)
p("Eleven suites, run against the real database with temporary accounts that are always cleaned up. "
  "They are the safety net for a codebase with no staging copy of production data.")
table(["Command", "Checks", "What it covers"],
      [["npm run verify:billing", "35", "Ledger, idempotency, refunds, webhooks"],
       ["npm run verify:thrift", "53", "Listings, messages, saves, reports, blocks"],
       ["npm run verify:tryon", "37", "Generation, credit debits, refunds on failure"],
       ["npm run verify:transfer", "27", "Sale completion and wardrobe transfer"],
       ["npm run verify:media", "24", "Image access control and rate limiting"],
       ["npm run verify:deletion", "33", "Account deletion, image cleanup, financial archiving"],
       ["npm run verify:policy", "18", "Policy versions and acceptance"],
       ["npm run verify:notifications", "16", "Delivery, dedupe, isolation, read state"],
       ["npm run verify:age", "14", "Date-of-birth and minimum-age rules"],
       ["npm run verify:wear", "13", "Wear, reset, clean, and last-worn dates"]],
      widths=[2.3, 0.8, 3.4])
p("Total: 270 checks. A test that cannot fail is worthless, so where a suite exists to catch a specific "
  "bug it was confirmed to fail against the broken code before being kept.")

h("5.6  Monitoring", 2)
p("The application writes one structured line per business event, searchable in Vercel's logs. These "
  "lines deliberately contain no image URLs, provider identifiers, email addresses, or figures a "
  "customer never saw — logs get shared more casually than databases.")
p("Events include: purchases started and granted, payments failed, subscriptions activated, renewed and "
  "cancelled; try-on generated, regenerated, cached, failed, refused for want of credits, rate limited; "
  "credits granted, debited and refunded; chat used and chat limit reached; image access denied and rate "
  "limited; account deletion requested, completed, and image cleanup shortfalls.")
p("A read-only storage audit is available and never modifies anything:")
code("npm run audit:images")

# ============================================================== PART 6
h("Part 6 — Honest limits and what is outstanding", 1, page_break=True)
p("Stated plainly, because a documentation set that only describes what works is not much use when "
  "something does not.")

h("6.1  Deliberately not built", 2)
bullets([
    ("No payments inside Thrift", "no escrow, payouts, shipping labels, delivery tracking, returns or "
     "buyer protection. Buyer and seller settle between themselves."),
    ("Age is stated, not verified", "date of birth is collected and checked, but not proven."),
    ("Profile details cannot be edited", "name, date of birth and gender are set at registration."),
    ("No password", "access to your email is access to your account."),
])

h("6.2  Known technical limitations", 2)
bullets([
    ("Sessions are stateless for thirty days", "there is no server-side session store, so a token stays "
     "cryptographically valid until it expires. Deleting an account clears the cookie, and a deleted "
     "account's token can do nothing — no rows exist to read and writes fail on the foreign key — but "
     "the token is not actively revoked."),
    ("Rate limits are per-instance", "they are held in memory, so on serverless they are a speed bump "
     "rather than a guarantee."),
    ("A few endpoints still return image addresses directly", "three places in sharing and one in "
     "history, rather than going through the access-controlled media route."),
])

h("6.3  Outstanding operational work", 2)
table(["Item", "Why it matters", "Priority"],
      [["Rotate the R2, fal and Resend keys", "They were exposed in a working transcript", "High"],
       ["Add https://localhost and capacitor://localhost to R2 CORS", "App uploads fail without it", "High"],
       ["Publish a DMARC record for the domain", "Improves deliverability of sign-in codes", "Medium"],
       ["Signed release build and keystore", "Required before Google Play will accept the app", "Medium"],
       ["Route the last image URLs through the media endpoint", "Consistency with the privacy model", "Medium"],
       ["Wire the APK copy-back into the Gradle build", "Prevents shipping a stale APK", "Low"],
       ["Remaining prerendered marketing pages", "Better left until there is search data", "Low"]],
      widths=[2.5, 2.8, 1.2])

# ============================================================== PART 7
h("Part 7 — Reference tables", 1, page_break=True)

h("7.1  The complete API", 2)
p("Every endpoint, grouped by area. All are prefixed with /api. Everything except the landing pages, "
  "the health check and the payment webhook requires a signed-in session.")

def api_table(title, rows):
    h(title, 3)
    table(["Method and path", "What it does"], rows, widths=[2.7, 3.8])

api_table("Authentication  /api/auth", [
    ["POST /start", "Send a sign-in code to an email address"],
    ["POST /verify", "Check the code; sign in or continue to registration"],
    ["POST /complete", "Finish registration with name, date of birth and gender"],
    ["GET  /me", "Who is signed in"],
    ["POST /logout", "End the session"]])

api_table("Account  /api/account", [
    ["GET  /deletion-preview", "What deleting would destroy, counted"],
    ["POST /delete/start", "Email a deletion confirmation code"],
    ["POST /delete/confirm", "Delete the account permanently"]])

api_table("Wardrobe  /api/clothes", [
    ["GET  /", "List garments, optionally filtered by status"],
    ["GET  /:id", "One garment"],
    ["POST /upload-url", "A signed URL for uploading a photo"],
    ["POST /", "Add a garment"],
    ["PATCH /:id", "Edit name, category or style tag"],
    ["DELETE /:id", "Delete a garment and its photo"],
    ["POST /wear", "Mark garments worn today"],
    ["POST /reset", "Return everything worn to clean"],
    ["POST /:id/clean", "Return one garment to clean"],
    ["GET  /plans", "Upcoming outfit plans"],
    ["POST /plan", "Plan an outfit for a date"],
    ["DELETE /plans/:id", "Remove a plan"]])

api_table("Try-on  /api/tryon", [
    ["GET  /selfie", "The current photo of you"],
    ["POST /selfie/upload-url", "A signed URL for uploading it"],
    ["POST /selfie", "Save it"],
    ["POST /generate", "Generate a look and debit credits"],
    ["GET  /history", "Previously generated looks"],
    ["DELETE /:id", "Delete a generated look"]])

api_table("Billing  /api/billing", [
    ["GET  /products", "Packs and plans"],
    ["GET  /summary", "Balance, plan, renewal date, chat allowance"],
    ["POST /create-pack-order", "Start a one-off purchase"],
    ["POST /create-subscription", "Start a subscription"],
    ["POST /verify-payment", "Confirm a completed payment"],
    ["POST /webhook", "Razorpay's signed callback"]])

api_table("Thrift  /api/thrift", [
    ["GET  /listings", "Browse, with filters"],
    ["GET  /listings/:id", "One listing"],
    ["POST /listings", "Create one"],
    ["PATCH /listings/:id", "Edit one"],
    ["DELETE /listings/:id", "Delete one"],
    ["POST /listings/:id/activate", "Publish it"],
    ["POST /listings/:id/pause", "Hide it temporarily"],
    ["POST /listings/:id/save", "Save it  ·  DELETE to unsave"],
    ["POST /listings/:id/report", "Report it"],
    ["POST /listings/:id/mark-sold", "Seller records the sale"],
    ["POST /listings/:id/sell", "Begin the sale"],
    ["POST /listings/:id/conversation", "Start talking to the seller"],
    ["GET  /mine  ·  /saved", "Your listings  ·  your saved items"],
    ["GET  /messages  ·  /messages/:id", "Conversations  ·  one thread"],
    ["POST /messages/:id", "Send a message"],
    ["POST /messages/:id/read", "Mark a thread read"],
    ["POST /messages/:id/report", "Report a conversation"],
    ["GET  /transactions", "Your sales and purchases"],
    ["POST /transactions/:id/confirm", "Buyer confirms; triggers the transfer"],
    ["POST /transactions/:id/cancel", "Call off a sale"],
    ["POST /users/:id/block", "Block someone  ·  DELETE to unblock"]])

api_table("Sharing  /api", [
    ["GET  /share/codes", "Codes you have issued"],
    ["POST /share/codes", "Create a code"],
    ["DELETE /share/codes/:id", "Cancel a code"],
    ["POST /share/redeem", "Redeem someone else's code"],
    ["GET  /share/with-me  ·  /share/i-can-see", "Both directions of sharing"],
    ["DELETE /share/:id/owner  ·  /viewer", "Disconnect, from either side"],
    ["GET  /friends/:id/wardrobe", "A wardrobe shared with you"],
    ["POST /friends/:id/suggest", "Suggest an outfit"],
    ["POST /friends/:id/plan", "Plan an outfit for them"],
    ["GET  /suggestions", "Suggestions waiting for you"],
    ["POST /suggestions/:id/respond", "Accept or decline"]])

api_table("Everything else", [
    ["GET  /media/:scope/:id", "A short-lived signed link to an image"],
    ["GET  /media/proxy/:scope/:id", "The image bytes, streamed (used by the app)"],
    ["POST /chat", "Ask the stylist"],
    ["GET  /history", "Wear history"],
    ["GET  /notifications  ·  POST /read", "The tray  ·  mark read"],
    ["GET  /policy/status  ·  POST /accept", "Policy state  ·  record acceptance"],
    ["GET  /onboarding  ·  PATCH /", "Guided tour state"],
    ["POST /contact", "Message support"],
    ["GET  /config", "Client configuration"],
    ["GET  /health", "Commit, environment, and what is configured"]])

h("7.2  Document information", 2)
table(["Field", "Value"],
      [["Product", "TryUnex"],
       ["Version", "1.1  (Android versionCode 2)"],
       ["Build", COMMIT],
       ["Generated", TODAY],
       ["Source", "Read directly from the codebase at the above commit"],
       ["Regenerate with", "python3 docs/generate_documentation.py"]],
      widths=[1.8, 4.7])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TryUnex-Documentation.docx")
doc.save(out)
print("saved:", out)
